#!/usr/bin/env python3
"""메모복붙 앱 사용설명서 DOCX 생성 (폴라리스 오피스 호환)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# 기본 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = 'Malgun Gothic'
font.size = Pt(11)

# ========== 표지 ==========
for _ in range(6):
    doc.add_paragraph('')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('메모복붙')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x21, 0x96, 0xF3)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('앱 사용설명서')
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph('')
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ver.add_run('버전 1.0.0  |  2026년 3월')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.add_page_break()

# ========== 목차 ==========
h = doc.add_heading('목차', level=1)
toc_items = [
    '1. 앱 소개',
    '2. 메인 화면 (홈)',
    '3. 카테고리 관리',
    '4. 스니펫 관리',
    '5. 변수 치환 기능',
    '6. 검색 기능',
    '7. 클립보드 히스토리',
    '8. 설정',
    '9. 자주 묻는 질문 (FAQ)',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(13)

doc.add_page_break()

# ========== 1. 앱 소개 ==========
doc.add_heading('1. 앱 소개', level=1)
doc.add_paragraph(
    '메모복붙은 자주 사용하는 문구(스니펫)를 카테고리별로 정리하고, '
    '한 번의 터치로 클립보드에 복사할 수 있는 앱입니다.'
)
doc.add_paragraph('')

doc.add_heading('주요 특징', level=2)
features = [
    ('카테고리 분류', '계좌번호, 주소, 이메일 템플릿 등 용도별 분류'),
    ('원터치 복사', '스니펫 옆 복사 버튼 한 번으로 클립보드 복사'),
    ('변수 치환', '{{고객이름}}, {{날짜}} 등 동적 변수 지원'),
    ('클립보드 히스토리', '복사한 내용을 자동으로 기록'),
    ('검색', '전체 스니펫에서 실시간 검색'),
    ('클라우드 동기화', 'Firebase를 통한 백업 및 동기화'),
    ('iOS 키보드 확장', '다른 앱에서 바로 스니펫 입력 가능'),
]
for title_text, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f'• {title_text}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== 2. 메인 화면 ==========
doc.add_heading('2. 메인 화면 (홈)', level=1)

doc.add_heading('화면 구성', level=2)
doc.add_paragraph(
    '앱을 실행하면 메인 화면이 표시됩니다. 상단에 통계 정보, '
    '중앙에 카테고리 그리드가 표시됩니다.'
)

doc.add_heading('상단 통계 영역', level=3)
stats = [
    ('전체 스니펫', '등록된 전체 스니펫 수'),
    ('총 복사', '지금까지 복사한 총 횟수'),
]
for label, desc in stats:
    p = doc.add_paragraph()
    run = p.add_run(f'• {label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('상단 우측 아이콘 버튼', level=3)
buttons = [
    ('시계 아이콘', '클립보드 히스토리 화면으로 이동'),
    ('돋보기 아이콘', '검색 화면으로 이동'),
    ('연필 아이콘', '편집 모드 (카테고리 삭제/수정 가능)'),
    ('톱니바퀴 아이콘', '설정 화면으로 이동'),
]
for label, desc in buttons:
    p = doc.add_paragraph()
    run = p.add_run(f'• {label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('카테고리 그리드', level=3)
doc.add_paragraph(
    '카테고리가 2열 그리드로 표시됩니다. 각 카테고리 카드에는 '
    '아이콘, 이름, 스니펫 개수가 표시됩니다.'
)
actions = [
    ('카테고리 탭', '해당 카테고리의 스니펫 목록으로 이동'),
    ('카테고리 길게 누르기', '카테고리 수정 다이얼로그 표시'),
    ('편집 모드에서 탭', '카테고리 수정'),
    ('"카테고리 추가" 카드 탭', '새 카테고리 생성'),
]
for label, desc in actions:
    p = doc.add_paragraph()
    run = p.add_run(f'• {label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== 3. 카테고리 관리 ==========
doc.add_heading('3. 카테고리 관리', level=1)

doc.add_heading('기본 카테고리', level=2)
doc.add_paragraph('앱 설치 시 4개의 기본 카테고리가 생성됩니다:')
defaults = ['일반 메모', '계좌번호', '주소', '이메일 템플릿']
for d in defaults:
    doc.add_paragraph(f'• {d}')

doc.add_heading('카테고리 추가', level=2)
steps = [
    '메인 화면에서 "카테고리 추가" 카드를 탭합니다.',
    '카테고리 이름을 입력합니다. (예: 인사말, 안내문구 등)',
    '12가지 색상 중 원하는 색상을 선택합니다.',
    '16가지 아이콘 중 원하는 아이콘을 선택합니다.',
    '"만들기" 버튼을 탭하면 카테고리가 생성됩니다.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('카테고리 수정', level=2)
doc.add_paragraph(
    '카테고리를 길게 누르거나, 편집 모드(연필 아이콘)에서 카테고리를 탭하면 '
    '수정 다이얼로그가 표시됩니다. 이름, 색상, 아이콘을 변경할 수 있습니다.'
)

doc.add_heading('카테고리 삭제', level=2)
doc.add_paragraph(
    '편집 모드에서 카테고리 카드 우측 상단의 빨간 X 버튼을 탭합니다. '
    '확인 다이얼로그에서 "삭제"를 선택하면 카테고리와 포함된 모든 스니펫이 삭제됩니다.'
)
p = doc.add_paragraph()
run = p.add_run('⚠ 주의: 삭제된 카테고리와 스니펫은 복구할 수 없습니다.')
run.bold = True
run.font.color.rgb = RGBColor(0xF4, 0x43, 0x36)

doc.add_page_break()

# ========== 4. 스니펫 관리 ==========
doc.add_heading('4. 스니펫 관리', level=1)

doc.add_heading('스니펫이란?', level=2)
doc.add_paragraph(
    '스니펫은 자주 사용하는 문구 조각입니다. '
    '계좌번호, 주소, 인사말, 이메일 본문 등 반복적으로 입력하는 텍스트를 '
    '미리 등록해두고 필요할 때 복사하여 사용합니다.'
)

doc.add_heading('스니펫 추가', level=2)
steps = [
    '카테고리를 탭하여 스니펫 목록 화면으로 이동합니다.',
    '우측 하단의 + 버튼을 탭합니다.',
    '스니펫 유형을 선택합니다 (일반, 계좌, 주소, 이메일, 코드).',
    '제목을 입력합니다. (선택사항 - 미입력 시 내용에서 자동 생성)',
    '내용을 입력합니다.',
    '필요 시 태그를 추가합니다. (#해시태그 형태)',
    '상단 "저장" 버튼을 탭합니다.',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('스니펫 유형', level=2)
types = [
    ('일반 (text)', '범용 텍스트 메모'),
    ('계좌 (account)', '은행 계좌번호'),
    ('주소 (address)', '주소, 위치 정보'),
    ('이메일 (email)', '이메일 템플릿'),
    ('코드 (code)', '코드 스니펫'),
]
for label, desc in types:
    p = doc.add_paragraph()
    run = p.add_run(f'• {label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('스니펫 복사', level=2)
doc.add_paragraph(
    '스니펫 목록에서 우측의 복사 버튼(📋)을 탭하면 '
    '클립보드에 즉시 복사됩니다. 복사 완료 시 버튼이 초록색 체크 표시로 변합니다.'
)
doc.add_paragraph(
    '변수가 포함된 스니펫의 경우, 복사 시 변수 입력 다이얼로그가 표시됩니다.'
)

doc.add_heading('스니펫 고정 (핀)', level=2)
doc.add_paragraph(
    '스니펫을 왼쪽으로 스와이프하면 "고정" 버튼이 나타납니다. '
    '고정된 스니펫은 목록 최상단에 표시되며, 📌 아이콘이 표시됩니다.'
)

doc.add_heading('스니펫 삭제', level=2)
doc.add_paragraph(
    '스니펫을 왼쪽으로 스와이프하면 "삭제" 버튼(빨간색)이 나타납니다. '
    '탭하면 스니펫이 삭제됩니다.'
)

doc.add_heading('스니펫 순서 변경', level=2)
doc.add_paragraph(
    '카테고리 상세 화면에서 상단의 순서 변경 아이콘을 탭하면 '
    '드래그 앤 드롭으로 스니펫 순서를 변경할 수 있습니다.'
)

doc.add_heading('스니펫 공유', level=2)
doc.add_paragraph(
    '스니펫 편집 화면에서 공유 버튼을 탭하면 '
    '다른 앱(카카오톡, 메시지 등)으로 내용을 공유할 수 있습니다.'
)

doc.add_page_break()

# ========== 5. 변수 치환 기능 ==========
doc.add_heading('5. 변수 치환 기능', level=1)
doc.add_paragraph(
    '메모복붙의 핵심 기능입니다. 스니펫에 {{변수명}} 형태로 변수를 넣으면, '
    '복사할 때 자동으로 치환되거나 입력을 요청합니다.'
)

doc.add_heading('내장 변수 (자동 치환)', level=2)
doc.add_paragraph('아래 변수는 복사 시 자동으로 현재 값으로 치환됩니다:')
builtin = [
    ('{{날짜}}', '오늘 날짜 (예: 2026-03-18)'),
    ('{{시간}}', '현재 시간 (예: 14:30)'),
    ('{{요일}}', '오늘 요일 (예: 수)'),
    ('{{#카운터}}', '복사 횟수 + 1 (예: 5)'),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = '변수'
hdr[1].text = '설명'
for cell in hdr:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True
for var, desc in builtin:
    row = table.add_row().cells
    row[0].text = var
    row[1].text = desc

doc.add_paragraph('')

doc.add_heading('사용자 변수 (입력 필요)', level=2)
doc.add_paragraph(
    '내장 변수 외의 모든 {{...}} 패턴은 사용자 변수입니다. '
    '복사 시 입력 다이얼로그가 나타나며, 각 변수에 값을 입력한 후 '
    '복사하면 치환된 결과가 클립보드에 복사됩니다.'
)

doc.add_heading('사용 예시', level=2)
doc.add_paragraph('스니펫 내용:')
p = doc.add_paragraph()
run = p.add_run(
    '{{고객이름}}님 안녕하세요.\n'
    '{{날짜}} 주문하신 상품이 발송되었습니다.\n'
    '배송지: {{배송주소}}\n'
    '감사합니다.'
)
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_paragraph('')
doc.add_paragraph('복사 시:')
items = [
    '{{날짜}} → 자동으로 오늘 날짜 치환',
    '{{고객이름}} → 입력 다이얼로그에서 입력',
    '{{배송주소}} → 입력 다이얼로그에서 입력',
]
for item in items:
    doc.add_paragraph(f'• {item}')

doc.add_page_break()

# ========== 6. 검색 기능 ==========
doc.add_heading('6. 검색 기능', level=1)
doc.add_paragraph(
    '메인 화면 상단의 돋보기 아이콘을 탭하면 검색 화면으로 이동합니다.'
)

search_features = [
    ('실시간 검색', '입력하는 즉시 결과가 업데이트됩니다'),
    ('전체 검색', '모든 카테고리의 스니펫에서 검색합니다'),
    ('검색 대상', '제목, 내용, 태그를 모두 검색합니다'),
    ('카테고리 표시', '검색 결과에 소속 카테고리가 표시됩니다'),
]
for label, desc in search_features:
    p = doc.add_paragraph()
    run = p.add_run(f'• {label}: ')
    run.bold = True
    p.add_run(desc)

# ========== 7. 클립보드 히스토리 ==========
doc.add_heading('7. 클립보드 히스토리', level=1)
doc.add_paragraph(
    '메인 화면 상단의 시계 아이콘을 탭하면 클립보드 히스토리 화면으로 이동합니다. '
    '앱에서 복사한 모든 내용이 자동으로 기록됩니다.'
)

history_features = [
    ('자동 기록', '스니펫 복사 시 자동으로 히스토리에 저장'),
    ('재복사', '히스토리 항목을 탭하면 다시 클립보드에 복사'),
    ('개별 삭제', '왼쪽 스와이프 후 삭제 버튼 탭'),
    ('전체 삭제', '상단 휴지통 아이콘으로 전체 히스토리 삭제'),
]
for label, desc in history_features:
    p = doc.add_paragraph()
    run = p.add_run(f'• {label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_page_break()

# ========== 8. 설정 ==========
doc.add_heading('8. 설정', level=1)

doc.add_heading('복사 설정', level=2)
settings = [
    ('복사 피드백', '복사 완료 시 확인 메시지 표시 여부'),
    ('내장 변수 자동 치환', '{{날짜}}, {{시간}}, {{요일}} 자동 치환 여부'),
]
for label, desc in settings:
    p = doc.add_paragraph()
    run = p.add_run(f'• {label}: ')
    run.bold = True
    p.add_run(desc)

doc.add_heading('클라우드 동기화', level=2)
doc.add_paragraph(
    'Firebase 계정을 통해 데이터를 클라우드에 백업하고 동기화할 수 있습니다.'
)
sync_steps = [
    '설정 화면에서 "로그인" 버튼을 탭합니다.',
    '이메일과 비밀번호를 입력하여 로그인 또는 회원가입합니다.',
    '"지금 동기화" 버튼으로 수동 동기화를 실행합니다.',
    '마지막 동기화 시간이 표시됩니다.',
]
for i, step in enumerate(sync_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('데이터 삭제', level=2)
p = doc.add_paragraph()
run = p.add_run('⚠ "전체 데이터 삭제"를 탭하면 모든 카테고리와 스니펫이 영구 삭제됩니다. ')
run.font.color.rgb = RGBColor(0xF4, 0x43, 0x36)
run.bold = True
doc.add_paragraph('이 작업은 되돌릴 수 없으므로 신중하게 사용하세요.')

doc.add_page_break()

# ========== 9. FAQ ==========
doc.add_heading('9. 자주 묻는 질문 (FAQ)', level=1)

faqs = [
    ('Q. 스니펫은 최대 몇 개까지 등록할 수 있나요?',
     'A. 제한 없이 등록할 수 있습니다. 기기 저장 용량만 충분하면 됩니다.'),
    ('Q. 카테고리는 최대 몇 개까지 만들 수 있나요?',
     'A. 제한 없이 만들 수 있습니다.'),
    ('Q. 변수는 하나의 스니펫에 여러 개 넣을 수 있나요?',
     'A. 네, 원하는 만큼 넣을 수 있습니다. 같은 변수명을 여러 번 사용하면 한 번만 입력해도 모두 치환됩니다.'),
    ('Q. 클립보드 히스토리는 자동으로 삭제되나요?',
     'A. 자동 삭제되지 않습니다. 수동으로 개별 또는 전체 삭제할 수 있습니다.'),
    ('Q. 다른 기기에서 같은 데이터를 사용할 수 있나요?',
     'A. 클라우드 동기화 기능으로 여러 기기에서 동일한 데이터를 사용할 수 있습니다.'),
    ('Q. iOS 키보드에서 메모복붙을 사용하려면?',
     'A. 설정 > 일반 > 키보드 > 키보드 추가 에서 "메모복붙"을 추가하세요. '
     '이후 키보드 전환 버튼(🌐)으로 메모복붙 키보드를 선택하면 바로 스니펫을 입력할 수 있습니다.'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)
    doc.add_paragraph('')

# ========== 저장 ==========
output_path = '/Users/sangdonlee/memo-copypaste/메모복붙_사용설명서.docx'
doc.save(output_path)
print(f'저장 완료: {output_path}')
