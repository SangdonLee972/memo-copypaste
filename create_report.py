from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── 기본 스타일 설정 ──
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# 여백 설정
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    return h

def add_para(text, bold=False, size=None, color=None, align=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = '맑은 고딕'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if bold:
        run.bold = True
    if size:
        run.font.size = size
    if color:
        run.font.color.rgb = color
    if align:
        p.alignment = align
    p.paragraph_format.space_after = space_after
    return p

def add_screenshot(path, width=Inches(2.5), caption=""):
    if not os.path.exists(path):
        add_para(f"[이미지 없음: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        r.font.name = '맑은 고딕'
        r.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        cap.paragraph_format.space_after = Pt(12)

def add_table_row(table, cells_data, bold=False, bg_color=None):
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.name = '맑은 고딕'
        run.font.size = Pt(9)
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if bold:
            run.bold = True
        if bg_color:
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): bg_color,
                qn('w:val'): 'clear'
            })
            shading.append(shading_elm)
    return row

screenshots_dir = "C:/Users/Administrator/memo_copypaste/screenshots"

# ═══════════════════════════════════════════
# 표지
# ═══════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

add_para("메모복붙 앱 개발", bold=True, size=Pt(28), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("진행경과 보고서", bold=True, size=Pt(22), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

add_para("─" * 40, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))

add_para("보고일: 2026년 3월 8일", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("플랫폼: Flutter (Android / iOS / Web)", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para("패키지명: com.copynote.memo_copypaste", size=Pt(12), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. 프로젝트 개요
# ═══════════════════════════════════════════
add_heading_styled("1. 프로젝트 개요", level=1)

add_para("'메모복붙'은 자주 사용하는 문구를 카테고리별로 관리하고, 원탭으로 클립보드에 복사할 수 있는 모바일 앱입니다. "
         "비즈니스 인사말, 계좌번호, 주소, 이메일 템플릿 등 반복적으로 사용하는 텍스트를 효율적으로 관리할 수 있습니다.",
         space_after=Pt(12))

add_para("주요 특징:", bold=True, space_after=Pt(4))
features = [
    "카테고리별 문구(스니펫) 분류 및 관리",
    "원탭 클립보드 복사",
    "변수 치환 템플릿 ({{고객이름}}, {{날짜}} 등 동적 값 입력)",
    "클립보드 복사 히스토리",
    "태그 기반 분류 및 전체 검색",
    "라이트/다크 모드 지원",
]
for f in features:
    add_para(f"  •  {f}", size=Pt(10), space_after=Pt(2))

doc.add_paragraph()

# ═══════════════════════════════════════════
# 2. 전체 진행률
# ═══════════════════════════════════════════
add_heading_styled("2. 전체 진행률", level=1)

add_para("약 85% 완료", bold=True, size=Pt(16), color=RGBColor(0x22, 0x7C, 0x9D), space_after=Pt(12))

# 진행률 테이블
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
for i, text in enumerate(["구분", "상태", "비고"]):
    hdr[i].text = text
    for p in hdr[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

progress_data = [
    ["핵심 기능 개발", "완료", "카테고리, 스니펫, 복사, 검색 등"],
    ["UI/UX 디자인", "완료", "라이트 테마 적용, 카드형 레이아웃"],
    ["데이터베이스", "완료", "SQLite(sqflite) 연동"],
    ["클립보드 히스토리", "완료", "복사 이력 관리 기능"],
    ["변수 치환 템플릿", "완료", "{{변수명}} 동적 입력"],
    ["버그 수정", "진행중", "경미한 이슈 4건 잔여"],
    ["QA 테스트", "예정", "전체 기능 통합 테스트"],
    ["스토어 출시 준비", "예정", "아이콘, 스크린샷, 설명문"],
]
for row_data in progress_data:
    add_table_row(table, row_data)

doc.add_paragraph()

# ═══════════════════════════════════════════
# 3. 완료된 기능 (스크린샷)
# ═══════════════════════════════════════════
add_heading_styled("3. 완료된 기능 상세 (에뮬레이터 캡처)", level=1)

# 3-1 홈 화면
add_heading_styled("3-1. 홈 화면", level=2)
add_para("카테고리를 그리드 형태로 한눈에 확인할 수 있는 메인 화면입니다. "
         "상단에 전체 스니펫 수와 총 복사 횟수 통계가 표시되며, "
         "각 카테고리는 커스텀 아이콘과 색상으로 구분됩니다.", space_after=Pt(8))

add_screenshot(f"{screenshots_dir}/home_check.png", Inches(2.8), "[홈 화면] 카테고리 그리드 + 통계 바")

# 3-2 카테고리 상세
add_heading_styled("3-2. 카테고리 상세 (스니펫 목록)", level=2)
add_para("카테고리를 탭하면 해당 카테고리의 스니펫 목록이 표시됩니다. "
         "각 스니펫은 제목, 내용 미리보기, 태그, 복사 횟수가 표시되며 "
         "우측 복사 버튼을 탭하면 즉시 클립보드에 복사됩니다.", space_after=Pt(8))

add_screenshot(f"{screenshots_dir}/02_category_detail.png", Inches(2.8), "[카테고리 상세] 스니펫 목록 - 원탭 복사, 태그, 변수 템플릿")

# 3-3 편집 모드
add_heading_styled("3-3. 편집 모드", level=2)
add_para("홈 화면에서 편집 아이콘을 탭하면 편집 모드로 전환됩니다. "
         "각 카테고리에 삭제(X) 버튼이 표시되고, 카테고리를 탭하면 "
         "이름, 색상, 아이콘을 변경할 수 있는 수정 다이얼로그가 나타납니다.", space_after=Pt(8))

# 편집모드 2장 나란히
t = doc.add_table(rows=1, cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cell1 = t.rows[0].cells[0]
cell2 = t.rows[0].cells[1]

p1 = cell1.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r1 = p1.add_run()
r1.add_picture(f"{screenshots_dir}/check5.png", width=Inches(2.4))

p2 = cell2.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run()
r2.add_picture(f"{screenshots_dir}/03_category_edit.png", width=Inches(2.4))

cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cap.add_run("[편집 모드] 카테고리 삭제 버튼(좌)  /  카테고리 수정 다이얼로그(우)")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
r.font.name = '맑은 고딕'
r.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc.add_paragraph()

# 3-4 클립보드 히스토리
add_heading_styled("3-4. 클립보드 히스토리", level=2)
add_para("복사한 문구의 이력이 자동으로 저장되어 이전에 복사한 내용을 "
         "다시 확인하고 재복사할 수 있습니다.", space_after=Pt(8))

add_screenshot(f"{screenshots_dir}/04_clipboard_history.png", Inches(2.8), "[클립보드 히스토리] 복사 이력 목록 + 재복사")

# 3-5 검색
add_heading_styled("3-5. 검색 화면", level=2)
add_para("전체 스니펫을 대상으로 키워드 검색이 가능합니다. "
         "제목과 내용에서 일치하는 스니펫을 찾아 바로 복사할 수 있습니다.", space_after=Pt(8))

add_screenshot(f"{screenshots_dir}/05_search.png", Inches(2.8), "[검색 화면] 전체 스니펫 검색")

doc.add_page_break()

# ═══════════════════════════════════════════
# 4. 잔여 이슈
# ═══════════════════════════════════════════
add_heading_styled("4. 잔여 이슈 및 수정 예정 사항", level=1)

add_para("현재 발견된 경미한 이슈 4건으로, 앱의 핵심 동작에는 영향이 없습니다.", space_after=Pt(8))

table2 = doc.add_table(rows=1, cols=4)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = table2.rows[0].cells
for i, text in enumerate(["#", "내용", "심각도", "상태"]):
    hdr2[i].text = text
    for p in hdr2[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

issues = [
    ["1", "검색 화면 내부 타입 처리 보완", "중", "수정 예정"],
    ["2", "커스텀 폰트(Pretendard) 전역 적용", "하", "수정 예정"],
    ["3", "설정 > 전체 데이터 삭제 기능 구현", "중", "수정 예정"],
    ["4", "Android 앱 표시 이름 한글화 (\"메모복붙\")", "하", "수정 예정"],
]
for row_data in issues:
    add_table_row(table2, row_data)

doc.add_paragraph()

# ═══════════════════════════════════════════
# 5. 향후 일정
# ═══════════════════════════════════════════
add_heading_styled("5. 향후 일정 (예상 1~2주)", level=1)

table3 = doc.add_table(rows=1, cols=4)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = table3.rows[0].cells
for i, text in enumerate(["주차", "단계", "내용", "예상 소요"]):
    hdr3[i].text = text
    for p in hdr3[i].paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.name = '맑은 고딕'
            run.font.size = Pt(9)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

schedule = [
    ["1주차", "버그 수정", "잔여 이슈 4건 수정 및 검증", "2~3일"],
    ["1주차", "QA 테스트", "전체 기능 통합 테스트, 다크모드, 엣지케이스 검증", "2~3일"],
    ["1주차", "UI 마무리", "앱 아이콘 제작, 폰트 적용, 세부 디자인 조정", "1~2일"],
    ["2주차", "출시 준비", "스토어용 스크린샷 제작, 앱 설명문 작성", "2~3일"],
    ["2주차", "최종 빌드", "릴리스 빌드 생성, 최종 테스트", "1~2일"],
    ["2주차", "스토어 등록", "Play Store / App Store 심사 제출", "1일"],
]
for row_data in schedule:
    add_table_row(table3, row_data)

doc.add_paragraph()

add_para("예상 완료일: 2026년 3월 21일 (±2~3일)", bold=True, size=Pt(11),
         color=RGBColor(0x22, 0x7C, 0x9D), space_after=Pt(16))

# ═══════════════════════════════════════════
# 6. 종합 평가
# ═══════════════════════════════════════════
add_heading_styled("6. 종합 평가", level=1)

add_para("앱의 핵심 기능인 카테고리 관리, 스니펫 원탭 복사, 변수 치환 템플릿, "
         "클립보드 히스토리, 편집 모드가 모두 구현 완료되어 에뮬레이터에서 정상 동작하는 것을 확인하였습니다.",
         space_after=Pt(8))

add_para("UI는 카드형 레이아웃으로 깔끔하게 구성되어 있으며, "
         "카테고리별 아이콘/컬러 커스터마이징, 태그 기반 분류, 복사 횟수 추적 등 "
         "사용자 편의 기능이 잘 구현되어 있습니다.",
         space_after=Pt(8))

add_para("남은 작업은 경미한 버그 수정과 출시 준비 단계로, "
         "코드 완성도는 높은 상태입니다. "
         "1~2주 내 최종 테스트 및 스토어 출시 준비를 완료할 예정입니다.",
         space_after=Pt(16))

add_para("─" * 40, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(8))
add_para("이상 보고 드립니다. 감사합니다.", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))

# 저장
output_path = "C:/Users/Administrator/memo_copypaste/메모복붙_진행경과보고서.docx"
doc.save(output_path)
print(f"보고서 생성 완료: {output_path}")
